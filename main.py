import os
import json
import re
import tempfile
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, Response, FileResponse
from pydantic import BaseModel
import fitz  # PyMuPDF

from dotenv import load_dotenv

from pymongo import MongoClient
from bson import ObjectId
from sentence_transformers import SentenceTransformer
import numpy as np

load_dotenv()  # <-- LOAD .env

PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")


if not PERPLEXITY_API_KEY:
    raise Exception("❌ Missing PERPLEXITY_API_KEY in .env")

# ===================== CONFIGURATION =====================
MONGO_URI = "mongodb+srv://aeranshorya_db_user:MnhnGxS22K8CzeGa@cluster0.figsico.mongodb.net/?appName=Cluster0"
DATABASE_NAME = "Coderower"
COLLECTION_NAME = "data"
EMBEDDING_MODEL = 'all-MiniLM-L6-v2'
OUTPUT_FOLDER = "filled_pdfs"
SIMILARITY_THRESHOLD = 0.70
KEYWORD_MATCH_RATIO = 0.30
MIN_KEYWORD_COUNT = 2


# ===================== INITIALIZE APP =====================
app = FastAPI(title="PDF RAG Form Filler", version="3.1")



app.mount("/filled_pdfs", StaticFiles(directory="filled_pdfs"), name="filled_pdfs")

# ===================== DATABASE CONNECTION =====================
try:
    client = MongoClient(MONGO_URI)
    db = client[DATABASE_NAME]
    knowledge_collection = db[COLLECTION_NAME]
    client.server_info()
    print("✅ MongoDB Atlas Connected Successfully")
except Exception as e:
    print(f"❌ MongoDB Connection Failed: {e}")

# ===================== MODEL LOADING =====================
try:
    model = SentenceTransformer(EMBEDDING_MODEL)
    print("✅ Embedding Model Loaded")
except Exception as e:
    print(f"❌ Model Loading Failed: {e}")

# ===================== LEARNING HELPERS =====================
def serialize_learning_record(record: Dict) -> Dict:
    """Convert MongoDB ObjectId to string for JSON serialization."""
    if record:
        record["id"] = str(record["_id"])
        record.pop("_id", None)
    return record

def create_learning_embedding(question: str, answer: str) -> List[float]:
    """Create embedding for learning record using question and answer."""
    try:
        combined_text = f"{question} {answer}"
        embedding = model.encode(combined_text).tolist()
        return embedding
    except Exception as e:
        print(f"❌ Embedding creation failed: {e}")
        return []

def clean_field_name(name: str) -> str:
    """Normalize PDF field labels."""
    import re
    name = re.sub(r"^\d+[\s._-]*", "", name)
    name = re.sub(r"[\s._-]*\d+$", "", name)
    return re.sub(r"\s+", " ", name).strip()

# ===================== VECTOR SEARCH =====================
def extract_filled_fields(pdf_path: str):
    result = []
    doc = fitz.open(pdf_path)

    for page in doc:
        widgets = page.widgets()
        if not widgets:
            continue

        for w in widgets:
            question = w.field_name
            answer = ""

            if w.field_value:
                answer = str(w.field_value).strip()

            result.append({
                "question": clean_field_name(question),
                "answer": answer,
                "raw_field": question
            })

    doc.close()
    return result

# # ===================== RAG SEARCH =====================
def find_answer(field_name: str, field_type: Optional[str] = None) -> Optional[str]:
    """
    FINAL CLEAN RAG (NO OPTIONAL SKIP):
    - No logs
    - Only returns the final answer string
    - Exact match + smart vector search
    - Keyword + similarity filters active
    """

    if knowledge_collection is None or model is None:
        return None

    cleaned = clean_field_name(field_name).lower()

    # --------------------------- EXACT MATCH ---------------------------
    try:
        exact = knowledge_collection.find_one(
            {"question": {"$regex": f"^{cleaned}$", "$options": "i"}}
        )
    except:
        exact = None

    if exact:
        ans = exact.get("answer", "")
        ok, _ = validate(cleaned, ans)
        if ok:
            return ans

    # --------------------------- VECTOR MATCH ---------------------------
    try:
        embed = model.encode(cleaned).tolist()

        pipeline = [
            {
                "$vectorSearch": {
                    "index": "default",
                    "path": "embedding",
                    "queryVector": embed,
                    "numCandidates": 40,
                    "limit": 5,
                }
            },
            {
                "$project": {
                    "question": 1,
                    "answer": 1,
                    "score": {"$meta": "vectorSearchScore"},
                }
            },
        ]

        results = list(knowledge_collection.aggregate(pipeline))
    except:
        return None

    if not results:
        return None

    # --------------------------- FILTERING ---------------------------
    for r in results:
        score = r.get("score", 0.0)
        if score < SIMILARITY_THRESHOLD:
            continue

        db_q = r.get("question", "")
        ans = r.get("answer", "")

        kw_ok, ratio, common = keyword_match(cleaned, db_q)
        if not kw_ok:
            continue

        ok, _ = validate(cleaned, ans)
        if not ok:
            continue

        return ans

    return None

# ===================== PDF HELPERS =====================
def extract_questions_from_pdf(pdf_path: str) -> List[Dict]:
    """Extract all form field names from a PDF."""
    questions = []
    try:
        doc = fitz.open(pdf_path)
        for page_num, page in enumerate(doc):
            widgets = page.widgets()
            if not widgets:
                continue
            for w in widgets:
                if not w.field_name:
                    continue
                questions.append({
                    "field_name": w.field_name,
                    "field_type": w.field_type_string,
                    "page": page_num + 1
                })
        doc.close()
        return questions
    except Exception as e:
        print(f"❌ PDF Extraction Error: {e}")
        return []

def extract_key_terms(text: str) -> set:
    stopwords = {"the", "a", "an", "and", "or", "in", "to", "for", "is", "are"}
    words = re.findall(r"\b\w+\b", text.lower())
    return {w for w in words if len(w) > 2 and w not in stopwords}


def is_date_format(x: str) -> bool:
    # Simple ISO-like date check, can extend later
    return bool(re.search(r"\d{4}-\d{2}-\d{2}", x))


def validate(field: str, answer: str):
    """Basic sanity checks so galat type ka data na bhare."""
    if not answer:
        return False, "empty"

    f = field.lower()
    a = answer.lower()

    if is_date_format(answer) and "date" not in f:
        return False, "date mismatch"

    if "sex" in f or "gender" in f:
        if a not in ["male", "female", "m", "f", "other"]:
            return False, "invalid gender"

    return True, "ok"


def keyword_match(field: str, db_question: str) -> Tuple[bool, float, set]:
    f_terms = extract_key_terms(field)
    q_terms = extract_key_terms(db_question)

    if not f_terms:
        return True, 1.0, set()

    common = f_terms & q_terms
    ratio = len(common) / len(f_terms) if f_terms else 0.0

    is_ok = (ratio >= KEYWORD_MATCH_RATIO) or (len(common) >= MIN_KEYWORD_COUNT)
    return is_ok, ratio, common

def fill_pdf_fields(pdf_path: str, answers: Dict[str, Optional[str]], output_path: str) -> Dict:
    """
    Sirf values fill karega, koi highlight nahi.
    Checkbox handling: ONLY matches exact button state; no hardcoded yes/no.
    """
    doc = fitz.open(pdf_path)
    filled = empty = checkboxes = 0

    for page in doc:
        widgets = page.widgets()
        if not widgets:
            continue

        for w in widgets:
            field = w.field_name
            ftype = w.field_type_string

            if field not in answers:
                continue

            ans = answers.get(field)
            if ans is None or str(ans).strip() == "":
                empty += 1
                continue

            ans = str(ans).strip()

            # Text fields
            if ftype not in ["Btn", "Ch"]:
                w.field_value = ans
                w.update()
                filled += 1
                continue

            # Checkbox / Radio
            ans_norm = ans.lower()
            try:
                states = w.button_states() or []
            except Exception:
                states = []

            matched = False
            for s in states:
                if not s:
                    continue
                if s.lower() == ans_norm:
                    w.field_value = s
                    w.update()
                    matched = True
                    filled += 1
                    checkboxes += 1
                    break

            if not matched:
                empty += 1

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()

    return {"filled": filled, "empty": empty, "checkboxes": checkboxes}


def add_red(page, rect: fitz.Rect):
    """Draw red box around given rect."""
    try:
        r = fitz.Rect(rect.x0 - 2, rect.y0 - 2, rect.x1 + 2, rect.y1 + 2)
        ann = page.add_rect_annot(r)
        ann.set_colors(stroke=(1, 0, 0))
        ann.set_border(width=2)
        ann.update()
    except Exception as e:
        print(f"Could not add red box: {e}")



def highlight_empty_groups(pdf_path: str, output_path: str) -> Dict:
    """
    Already-filled PDF ko scan karega:
    - Text + checkbox values dekhega
    - Group key: field_name ka prefix (split(' - ')[0])
    - Agar group me ek bhi option/field filled hai → us group pe red box NAHI
    - Agar group completely empty hai → us group ke combined rect pe red box
    """
    doc = fitz.open(pdf_path)

    # 1st pass: group info build karo
    groups: Dict[str, Dict] = {}
    for page_index, page in enumerate(doc):
        widgets = page.widgets()
        if not widgets:
            continue

        for w in widgets:
            field = w.field_name
            if not field:
                continue

            group_key = field.split(" - ")[0].strip()
            rect = w.rect
            ftype = w.field_type_string

            # Filled check:
            is_filled = False
            val = w.field_value

            if ftype in ["Btn", "Ch"]:
                # Most PDFs: "Off" = not checked
                if val not in (None, "", "Off"):
                    is_filled = True
            else:
                if val is not None and str(val).strip() != "":
                    is_filled = True

            if group_key not in groups:
                groups[group_key] = {
                    "page_index": page_index,
                    "rect": fitz.Rect(rect),
                    "filled": is_filled,
                }
            else:
                # Merge rects to cover whole area of that group
                groups[group_key]["rect"] = groups[group_key]["rect"] | rect
                groups[group_key]["filled"] = groups[group_key]["filled"] or is_filled

    # 2nd pass: sirf un groups pe red box jinke filled == False
    marked = 0
    for gkey, info in groups.items():
        if not info["filled"]:
            page = doc[info["page_index"]]
            add_red(page, info["rect"])
            marked += 1

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()

    return {"marked_groups": marked}

def extract_full_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])
    doc.close()
    return text

def count_all_fields(obj):
    if isinstance(obj, dict):
        total = 0
        for k, v in obj.items():
            if isinstance(v, dict):
                total += count_all_fields(v)
            else:
                total += 1
        return total
    return 0

import requests

def gpt_extract(raw_text):

    prompt = f"""
You are an intelligent information extraction engine.
Extract ALL fields and all answers from this PDF text.

Return ONLY JSON. No explanations. No markdown fences.

Document text:
--------------------------------
{raw_text}
--------------------------------
"""

    url = "https://api.perplexity.ai/chat/completions"

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}"
    }

    data = {
        "model": "sonar-pro",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0
    }

    response = requests.post(url, headers=headers, json=data)

    if response.status_code != 200:
        raise Exception(f"Perplexity API Error: {response.text}")

    content = response.json()["choices"][0]["message"]["content"]

    print("\n\n----- RAW PERPLEXITY OUTPUT -----\n")
    print(content)
    print("\n---------------------------------\n")

    # Remove markdown fences like ```json ... ```
    if content.startswith("```"):
        content = content.strip("`")
        if content.lower().startswith("json"):
            content = content[4:].strip()

    # Find the JSON block
    start = content.find("{")
    end = content.rfind("}")

    if start == -1 or end == -1:
        raise Exception(
            "Perplexity did not return JSON. Check RAW OUTPUT above."
        )

    json_text = content[start:end+1]

    try:
        return json.loads(json_text)
    except Exception as e:
        raise Exception(
            f"Could not parse JSON. Returned text:\n{json_text}\nError: {e}"
        )


def perplexity_call(prompt):
    """Generic Perplexity sonar-pro call that returns valid JSON."""
    url = "https://api.perplexity.ai/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}"
    }
    payload = {
        "model": "sonar-pro",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0
    }

    resp = requests.post(url, headers=headers, json=payload)
    if resp.status_code != 200:
        raise Exception(resp.text)

    content = resp.json()["choices"][0]["message"]["content"]

    content = content.replace("```json", "").replace("```", "").strip()

    # extract JSON list
    start = content.find("[")
    end = content.rfind("]")
    if start == -1 or end == -1:
        raise Exception("Invalid JSON: " + content)

    return json.loads(content[start:end + 1])



from docx import Document

def extract_text_from_docx(path):
    doc = Document(path)
    blocks = []

    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    blocks.append(txt)

    return "\n".join(blocks)


def extract_full_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])
    doc.close()
    return text

def pdf_has_real_fields(pdf_path):
    doc = fitz.open(pdf_path)
    for page in doc:
        widgets = page.widgets()
        if widgets:
            for w in widgets:
                if w.field_name:
                    doc.close()
                    return True
    doc.close()
    return False

@app.post("/api/extract-data/")
async def extract_dynamic(file: UploadFile = File(...)):

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Upload a PDF")

    # save temp PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(await file.read())
        pdf_path = tmp.name

    raw_text = extract_full_text(pdf_path)

    if not raw_text.strip():
        raise HTTPException(400, "PDF contains no text")

    json_data = gpt_extract(raw_text)

    return {
        "success": True,
        "fields_detected": count_all_fields(json_data),
        "data": json_data
    }

@app.post("/api/extract-questions/")
async def extract_questions(file: UploadFile = File(...)):
    """Extract field names from a PDF."""
    pdf_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(await file.read())
            pdf_path = tmp.name
        fields = extract_questions_from_pdf(pdf_path)
        return {"success": True, "fields": fields, "count": len(fields)}
    finally:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)

@app.post("/api/fill-pdf-with-rag/")
async def fill_pdf(file: UploadFile = File(...)):
    if knowledge_collection is None or model is None:
        raise HTTPException(400, "System not ready (DB/model missing)")

    # Save temp original PDF
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
    os.close(tmp_fd)
    with open(tmp_path, "wb") as f:
        f.write(await file.read())

    try:
        # 1) Extract all fields
        questions = extract_questions_from_pdf(tmp_path)

        answers: Dict[str, Optional[str]] = {}
        logs = []

        # 2) RAG answer per field
        for q in questions:
            name = q["field_name"]
            ftype = q["field_type"]

            res = find_answer(name, ftype)

            if res:
                answers[name] = res
                logs.append(
                    {
                        "field": name,
                        "answer": res,
                        "status": "filled"
                        # "method": res.get("method", "unknown"),
                        # "confidence": float(res.get("confidence", 0.0)),
                    }
                )
            else:
                answers[name] = None
                logs.append(
                    {
                        "field": name,
                        "answer": None,
                        "status": "empty",
                    }
                )

        # 3) Filled PDF (no highlights)
        filled_path = os.path.join(OUTPUT_FOLDER, f"filled_{file.filename}")
        filled_stats = fill_pdf_fields(tmp_path, answers, filled_path)

        # 4) Highlighted PDF (group-level red boxes for completely empty groups)
        highlighted_path = os.path.join(OUTPUT_FOLDER, f"highlighted_{file.filename}")
        highlight_stats = highlight_empty_groups(filled_path, highlighted_path)

        return {
            "success": True,
            "message": "✅ PDF filled successfully using MongoDB Vector Search.",
            "output": filled_path.replace("\\", "/"),
            "highlightedOutput": highlighted_path.replace("\\", "/"),
            "stats": {
                "filled": filled_stats["filled"],
                "empty": filled_stats["empty"],
                "checkboxes": filled_stats["checkboxes"],
                "marked_groups": highlight_stats["marked_groups"],
            },
            "logs": logs[:15],
        }

    except Exception as e:
        print("Error in /fill-pdf-with-rag/:", e)
        raise HTTPException(500, str(e))
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

@app.post("/api/learn-from-pdf/")
async def learn_from_pdf(
    category: str = Form(None),
    file: UploadFile = File(...)
):
    filename = file.filename.lower()
    ext = filename.split(".")[-1]

    # DB ready check
    if knowledge_collection is None or model is None:
        return {"success": False, "error": "Database or embedding model unavailable."}

    # Save file
    tmp_path = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}").name
    with open(tmp_path, "wb") as f:
        f.write(await file.read())

    now = datetime.utcnow().isoformat() + "Z"
    records = []

    try:
        # ============================================================
        # 1️⃣ DYNAMIC PDF → extract real form fields (NO Perplexity)
        # ============================================================
        if ext == "pdf" and pdf_has_real_fields(tmp_path):
            filled = extract_filled_fields(tmp_path)

            for row in filled:
                q = row["question"].strip()
                a = row["answer"].strip()

                if not a or a.lower() == "off":
                    continue

                embedding = model.encode(q).tolist()
                knowledge_collection.update_one(
                    {"question": q},
                    {
                        "$set": {
                            "question": q,
                            "answer": a,
                            "category": category,
                            "embedding": embedding,
                            "updatedAt": now
                        },
                        "$setOnInsert": {"createdAt": now, "createdBy": "dynamic-pdf"}
                    },
                    upsert=True
                )
                records.append({"question": q, "answer": a})

            return {
                "success": True,
                "mode": "dynamic-pdf",
                "inserted_or_updated": len(records),
                "records": records
            }

        # ============================================================
        # 2️⃣ DOCX → Perplexity extract Q/A JSON
        # ============================================================
        if ext == "docx":
            raw_text = extract_text_from_docx(tmp_path)

            prompt = f"""
Extract ALL question-answer pairs from this DOCX text.
Return ONLY pure JSON:
[
  {{"question": "...", "answer": "..."}}
]
TEXT:
{raw_text}
"""
            qa_list = perplexity_call(prompt)

        # ============================================================
        # 3️⃣ STATIC PDF → Perplexity extract Q/A JSON
        # ============================================================
        elif ext == "pdf":
            raw_text = extract_full_text(tmp_path)

            prompt = f"""
Extract ALL question-answer pairs from this PDF text.
Return ONLY JSON:
[
  {{"question": "...", "answer": "..."}}
]
TEXT:
{raw_text}
"""
            qa_list = perplexity_call(prompt)

        # ============================================================
        # Store Perplexity-extracted Q/A in DB
        # ============================================================
        for item in qa_list:
            q = item["question"].strip()
            a = item["answer"].strip()

            if not q or not a:
                continue

            embedding = model.encode(q).tolist()

            knowledge_collection.update_one(
                {"question": q},
                {
                    "$set": {
                        "question": q,
                        "answer": a,
                        "category": category,
                        "embedding": embedding,
                        "updatedAt": now
                    },
                    "$setOnInsert": {"createdAt": now, "createdBy": "perplexity"}
                },
                upsert=True
            )

            records.append({"question": q, "answer": a})

        return {
            "success": True,
            "mode": "perplexity",
            "inserted_or_updated": len(records),
            "records": records
        }

    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

# ===================== STARTUP EVENT =====================
@app.on_event("startup")
async def check_db():
    try:
        knowledge_count = knowledge_collection.count_documents({})
        print(f"✅ Startup: MongoDB connected")
        print(f"   📊 Knowledge collection: {knowledge_count} entries")
    except Exception as e:
        print(f"❌ Startup error: {e}")
